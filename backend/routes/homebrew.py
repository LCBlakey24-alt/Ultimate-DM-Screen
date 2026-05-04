"""Homebrew Workshop — upload a .docx (or paste text) and have an AI parse it
into a structured race / class / subclass / background / magic-item draft that
the user can then edit and save into their own personal ruleset.

Saves into the existing user_races / user_classes / user_subclasses /
user_backgrounds collections (so the Character Builder picks them up via the
existing /api/user/content endpoints), plus a new `user_magic_items` collection.

Uses Claude Sonnet via the configured AI provider for structured extraction.
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from pydantic import BaseModel
from typing import Optional, Dict, Any, List
import io
import json
import re
import uuid
from datetime import datetime, timezone

from utils.auth import get_current_user
from config import db, logger
from utils.llm_provider import LlmChat, UserMessage, get_llm_api_key

try:
    from docx import Document
except ImportError:
    Document = None

router = APIRouter()

CONTENT_TYPES = {"race", "class", "subclass", "background", "magic_item"}

# Mongo collections per content type
COLLECTION = {
    "race": "user_races",
    "class": "user_classes",
    "subclass": "user_subclasses",
    "background": "user_backgrounds",
    "magic_item": "user_magic_items",
}

# JSON schema hints for the LLM — kept compact for token efficiency
SCHEMA_HINTS = {
    "race": {
        "name": "string",
        "description": "string",
        "size": "Tiny|Small|Medium|Large",
        "speed": "int (default 30)",
        "ability_bonuses": "object e.g. {strength: 1, dexterity: 2}",
        "traits": "[{name, description}]",
        "languages": "[string]",
        "subraces": "[{name, description, ability_bonuses, traits}]",
    },
    "class": {
        "name": "string",
        "description": "string",
        "hit_die": "d6|d8|d10|d12",
        "primary_ability": "string",
        "saving_throw_proficiencies": "[string] — two abilities",
        "armor_proficiencies": "[string]",
        "weapon_proficiencies": "[string]",
        "features": "[{level: int 1-20, name, description}]",
    },
    "subclass": {
        "name": "string",
        "parent_class": "string — must match a class name",
        "description": "string",
        "subclass_level": "int (level the subclass unlocks, usually 3)",
        "features": "[{level: int 3-20, name, description}]",
    },
    "background": {
        "name": "string",
        "description": "string",
        "skill_proficiencies": "[string] — two skills",
        "tool_proficiencies": "[string]",
        "languages": "int — number of additional languages",
        "equipment": "[string]",
        "feature_name": "string",
        "feature_description": "string",
    },
    "magic_item": {
        "name": "string",
        "type": "Weapon|Armor|Wondrous Item|Potion|Ring|Rod|Scroll|Staff|Wand",
        "rarity": "common|uncommon|rare|very rare|legendary|artifact",
        "requires_attunement": "bool",
        "description": "string",
        "effects": "[string] — short bullet effects",
    },
}


def _docx_to_text(file_bytes: bytes) -> str:
    """Extract paragraph + table text from a .docx file."""
    if Document is None:
        raise HTTPException(status_code=503, detail="python-docx not installed on server")
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read .docx file: {type(e).__name__}")
    parts: List[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))
    return "\n".join(parts)


def _extract_json(reply: str) -> Optional[Dict[str, Any]]:
    """Pull the first JSON object from an LLM response (handles ```json fences)."""
    if not reply:
        return None
    # Strip code fences
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", reply, re.DOTALL)
    candidate = fenced.group(1) if fenced else None
    if candidate is None:
        # Fallback: find the largest balanced {...} substring
        first = reply.find("{")
        last = reply.rfind("}")
        if first != -1 and last > first:
            candidate = reply[first:last + 1]
    if candidate is None:
        return None
    try:
        return json.loads(candidate)
    except json.JSONDecodeError:
        return None


def _required_fields_for(content_type: str) -> List[str]:
    return {
        "race": ["name", "size", "speed"],
        "class": ["name", "hit_die", "features"],
        "subclass": ["name", "parent_class", "features"],
        "background": ["name", "skill_proficiencies"],
        "magic_item": ["name", "rarity"],
    }[content_type]


def _flag_missing(content_type: str, parsed: Dict[str, Any]) -> List[str]:
    """Return a list of required fields the AI couldn't fill in confidently."""
    missing: List[str] = []
    for field in _required_fields_for(content_type):
        value = parsed.get(field)
        if value in (None, "", [], {}):
            missing.append(field)
    return missing


async def _llm_extract(content_type: str, raw_text: str, username: str) -> Dict[str, Any]:
    """Ask Claude Sonnet to convert raw text into a structured JSON object
    matching SCHEMA_HINTS[content_type]. Returns a dict (possibly partial)."""
    if not LlmChat or not UserMessage or not get_llm_api_key("anthropic"):
        raise HTTPException(status_code=503, detail="AI is not configured on this server.")

    schema = SCHEMA_HINTS.get(content_type)
    if not schema:
        raise HTTPException(status_code=400, detail=f"Unsupported content_type '{content_type}'")

    schema_str = json.dumps(schema, indent=2)
    # Cap raw text to ~12k chars to keep tokens sane
    snippet = raw_text[:12000]

    system = (
        "You are an SRD-compliant TTRPG content extractor. Read the user's homebrew text and "
        "return ONLY a single JSON object that matches the requested schema. "
        "Use null or empty array/string for fields you cannot find. "
        "Do not include any explanation, only the JSON object. "
        "If the text describes class features at specific levels, include them in the features array."
    )
    prompt = (
        f"Content type: {content_type}\n"
        f"Schema (return JSON matching this shape):\n{schema_str}\n\n"
        f"Source text:\n{snippet}\n\n"
        f"Return ONLY the JSON object."
    )
    chat = LlmChat(
        api_key=get_llm_api_key("anthropic"),
        session_id=f"homebrew-{username}-{uuid.uuid4().hex[:6]}",
        system_message=system
    ).with_model("anthropic", "claude-sonnet-4-5-20250929")
    try:
        reply = await chat.send_message(UserMessage(text=prompt))
    except Exception as e:
        logger.exception("Homebrew LLM extract failed")
        raise HTTPException(status_code=502, detail=f"AI extraction failed: {type(e).__name__}")

    parsed = _extract_json(reply or "")
    if parsed is None:
        # Return an empty draft so the user can hand-fill rather than blocking
        return {}
    return parsed


# ─────────────────────────────────────────────────────────────────────────────
# Endpoints
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/homebrew/parse-docx")
async def parse_docx(
    content_type: str = Form(...),
    file: UploadFile = File(...),
    edition: str = Form("2014"),
    username: str = Depends(get_current_user)
):
    """Upload a .docx; we extract the text, ask Claude Sonnet to structure it
    into the right schema, and return a draft + a list of fields the AI
    couldn't fill in confidently."""
    if content_type not in CONTENT_TYPES:
        raise HTTPException(status_code=400, detail=f"content_type must be one of {sorted(CONTENT_TYPES)}")
    if edition not in ("2014", "2024"):
        edition = "2014"

    filename = (file.filename or "").lower()
    raw_bytes = await file.read()
    if not raw_bytes:
        raise HTTPException(status_code=400, detail="Empty file")

    if filename.endswith(".docx"):
        text = _docx_to_text(raw_bytes)
    elif filename.endswith(".txt") or filename.endswith(".md"):
        try:
            text = raw_bytes.decode("utf-8", errors="ignore")
        except Exception:
            raise HTTPException(status_code=400, detail="Could not read text file")
    else:
        raise HTTPException(status_code=400, detail="Only .docx, .txt, or .md files are supported")

    if not text.strip():
        raise HTTPException(status_code=400, detail="The file appears to be empty.")

    draft = await _llm_extract(content_type, text, username)
    missing = _flag_missing(content_type, draft)
    return {
        "content_type": content_type,
        "edition": edition,
        "draft": draft,
        "missing_fields": missing,
        "source_filename": file.filename,
        "source_excerpt": text[:1500]
    }


class ParseTextRequest(BaseModel):
    content_type: str
    edition: str = "2014"
    text: str


@router.post("/homebrew/parse-text")
async def parse_text(req: ParseTextRequest, username: str = Depends(get_current_user)):
    """Same as parse-docx but accepts pasted text instead of a file upload."""
    if req.content_type not in CONTENT_TYPES:
        raise HTTPException(status_code=400, detail=f"content_type must be one of {sorted(CONTENT_TYPES)}")
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="Text is empty")
    draft = await _llm_extract(req.content_type, req.text, username)
    missing = _flag_missing(req.content_type, draft)
    return {
        "content_type": req.content_type,
        "edition": req.edition if req.edition in ("2014", "2024") else "2014",
        "draft": draft,
        "missing_fields": missing,
        "source_excerpt": req.text[:1500]
    }


class HomebrewSaveRequest(BaseModel):
    content_type: str
    edition: str = "2014"
    data: Dict[str, Any]
    ruleset_id: Optional[str] = None  # we'll create/use a per-user "homebrew" ruleset
    homebrew_id: Optional[str] = None  # if set, we update; if not, we create


@router.post("/homebrew/save")
async def save_homebrew(req: HomebrewSaveRequest, username: str = Depends(get_current_user)):
    """Persist (or update) a homebrew item into the right user_* collection."""
    if req.content_type not in CONTENT_TYPES:
        raise HTTPException(status_code=400, detail="invalid content_type")
    coll_name = COLLECTION[req.content_type]

    # Find or create the user's "homebrew" ruleset bucket
    ruleset_id = req.ruleset_id
    if not ruleset_id:
        existing = await db.user_rulesets.find_one(
            {"user_id": username, "name": "Homebrew Workshop", "edition": req.edition},
            {"_id": 0}
        )
        if existing:
            ruleset_id = existing["id"]
        else:
            ruleset_id = str(uuid.uuid4())
            await db.user_rulesets.insert_one({
                "id": ruleset_id,
                "user_id": username,
                "name": "Homebrew Workshop",
                "description": "AI-assisted homebrew content",
                "edition": req.edition,
                "version": "1.0",
                "is_active": True,
                "created_at": datetime.now(timezone.utc).isoformat()
            })

    # Build doc
    doc = {**req.data}
    doc.update({
        "user_id": username,
        "ruleset_id": ruleset_id,
        "edition": req.edition,
        "source": "Homebrew Workshop",
        "updated_at": datetime.now(timezone.utc).isoformat()
    })

    if req.homebrew_id:
        # Update existing
        existing = await db[coll_name].find_one({"id": req.homebrew_id, "user_id": username})
        if not existing:
            raise HTTPException(status_code=404, detail="Homebrew item not found")
        await db[coll_name].update_one(
            {"id": req.homebrew_id, "user_id": username},
            {"$set": {k: v for k, v in doc.items() if k != "_id"}}
        )
        doc["id"] = req.homebrew_id
    else:
        doc["id"] = str(uuid.uuid4())
        doc["created_at"] = datetime.now(timezone.utc).isoformat()
        await db[coll_name].insert_one(doc)

    doc.pop("_id", None)
    return {"saved": True, "content_type": req.content_type, "homebrew": doc}


@router.get("/homebrew")
async def list_homebrew(
    content_type: Optional[str] = None,
    edition: Optional[str] = None,
    username: str = Depends(get_current_user)
):
    """List the current user's homebrew. Filterable by content_type and edition."""
    types = [content_type] if content_type else list(CONTENT_TYPES)
    out: Dict[str, List[Dict[str, Any]]] = {}
    for t in types:
        if t not in COLLECTION:
            continue
        q: Dict[str, Any] = {"user_id": username}
        if edition:
            q["edition"] = edition
        cursor = db[COLLECTION[t]].find(q, {"_id": 0})
        out[t] = [item async for item in cursor]
    return {"homebrew": out}


@router.delete("/homebrew/{content_type}/{homebrew_id}")
async def delete_homebrew(content_type: str, homebrew_id: str, username: str = Depends(get_current_user)):
    if content_type not in CONTENT_TYPES:
        raise HTTPException(status_code=400, detail="invalid content_type")
    coll_name = COLLECTION[content_type]
    result = await db[coll_name].delete_one({"id": homebrew_id, "user_id": username})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Homebrew item not found")
    return {"deleted": homebrew_id}
